import yaml
import pandas as pd
from docx import Document

# YAMLファイルのパス
yaml_file_path = 'config/yolov3_custom.yaml'

# YAMLファイルを読み込む
with open(yaml_file_path, 'r') as file:
    config = yaml.safe_load(file)

# データを整理してDataFrameに変換
data = {
    "Section": [],
    "Parameter": [],
    "Value": []
}

# Modelセクション
for key, value in config['model'].items():
    data["Section"].append("Model")
    data["Parameter"].append(key)
    data["Value"].append(value)

# Trainセクション
for key, value in config['train'].items():
    data["Section"].append("Train")
    data["Parameter"].append(key)
    data["Value"].append(value)

# Augmentationセクション
for key, value in config['augmentation'].items():
    data["Section"].append("Augmentation")
    data["Parameter"].append(key)
    data["Value"].append(value)

# Testセクション
for key, value in config['test'].items():
    data["Section"].append("Test")
    data["Parameter"].append(key)
    data["Value"].append(value)

# DataFrameの作成
df = pd.DataFrame(data)

# Wordドキュメントの作成
doc = Document()
doc.add_heading('YOLOv3 Configuration', level=1)

# 表の追加
table = doc.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Section'
hdr_cells[1].text = 'Parameter'
hdr_cells[2].text = 'Value'

# DataFrameの内容を表に追加
for index, row in df.iterrows():
    row_cells = table.add_row().cells
    row_cells[0].text = str(row['Section'])
    row_cells[1].text = str(row['Parameter'])
    row_cells[2].text = str(row['Value'])

# Wordファイルの保存
doc.save('yolov3_configuration.docx')